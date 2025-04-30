from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
import pandas as pd
import os
from docx.oxml import OxmlElement

def add_page_number(document):
    # Agregar número de página en la parte inferior derecha
    section = document.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"
    run._element.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar)

def main():
    # Crear el documento principal
    main_doc = Document()
    
    # Configurar el documento
    for section in main_doc.sections:
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
    
    # Agregar numeración de página
    add_page_number(main_doc)
    
    # Crear la portada
    para = main_doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture('portada.jpg', width=Inches(6))  # Ajusta la ruta y tamaño según sea necesario
    
    para = main_doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("COMPILACIÓN DE CUENTOS")
    run.font.size = Pt(24)
    run.font.bold = True
    
    # Agregar salto de sección
    main_doc.add_section(WD_SECTION.NEW_PAGE)
    
    # Página con frase motivadora
    para = main_doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run("A veces, una historia puede ser el puente que necesita un corazón para sanar.")
    run.font.size = Pt(16)
    run.font.bold = True
    
    para = main_doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    mensaje = """Este libro nace del alma de jóvenes valientes, que se atrevieron a compartir sus emociones, sueños, heridas y esperanzas. Cada cuento aquí escrito es un reflejo auténtico de sus caminos, de sus luchas y de su poder de transformación.
Que estas páginas sean no solo un testimonio, sino también una luz para quienes necesitan saber que no están solos.
Gracias a cada autor y autora por prestarnos su voz. Gracias por ser una inspiración...
Mayo de 2015."""
    para.add_run(mensaje)
    
    # Agregar salto de sección
    main_doc.add_section(WD_SECTION.NEW_PAGE)
    
    # Página con índice
    para = main_doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("ÍNDICE")
    run.font.size = Pt(16)
    run.font.bold = True
    
    # Cargar datos de Excel
    df = pd.read_excel('Cuentos por porgrama.xlsx')
    
    # Placeholder para el índice - lo completaremos al final
    index_placeholder = main_doc.add_paragraph()
    
    # Agregar salto de sección
    main_doc.add_section(WD_SECTION.NEW_PAGE)
    
    # Lista para almacenar información para el índice
    index_entries = []
    
    # Directorio con los archivos Word
    docx_dir = 'cuentos/'  # Ajusta según tu estructura
    
    # Procesar cada documento Word
    for filename in os.listdir(docx_dir):
        if filename.endswith('.docx') and not filename.startswith('~$'):
            cuento_path = os.path.join(docx_dir, filename)
            title = filename.replace('.docx', '')
            
            # Buscar información en Excel
            cuento_info = df[df['TITULO'].str.contains(title, case=False, na=False)]
            
            if cuento_info.empty:
                autor = "Autor Desconocido"
                programa = "Programa No Especificado"
            else:
                autor = cuento_info.iloc[0]['AUTOR']
                programa = cuento_info.iloc[0]['PROGRAMA']
            
            # Guardar para el índice
            index_entries.append(title)
            
            # Agregar el contenido del documento
            cuento_doc = Document(cuento_path)
            
            # Agregar cada párrafo del cuento
            for paragraph in cuento_doc.paragraphs:
                new_para = main_doc.add_paragraph()
                new_para.alignment = paragraph.alignment
                
                for run in paragraph.runs:
                    new_run = new_para.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    if run.font.size:
                        new_run.font.size = run.font.size
            
            # Agregar línea divisoria
            main_doc.add_paragraph().runs[0].add_break()
            border_para = main_doc.add_paragraph()
            border_para.paragraph_format.bottom_border.width = 1
            
            # Agregar información del cuento
            info_para = main_doc.add_paragraph()
            info_para.add_run(f"{title}\n").bold = True
            info_para.add_run(f"{autor}\n")
            info_para.add_run(f"{programa}")
            
            # Agregar salto de sección para el siguiente cuento
            main_doc.add_section(WD_SECTION.NEW_PAGE)
    
    # Completar el índice
    index_entries.sort()  # Ordenar alfabéticamente
    for i, title in enumerate(index_entries):
        index_placeholder.add_run(f"{title}\n")
    
    # Guardar el documento Word
    main_doc.save('Compilacion_Cuentos.docx')
    
    # Convertir a HTML (versión simplificada - para una conversión mejor se necesitarían herramientas adicionales)
    # Esto es solo un marcador de posición, la conversión real requeriría herramientas específicas
    with open('Compilacion_Cuentos.html', 'w', encoding='utf-8') as f:
        f.write('<!DOCTYPE html>\n<html>\n<head>\n<title>Compilación de Cuentos</title>\n')
        f.write('<meta charset="UTF-8">\n')
        f.write('<style>\n')
        f.write('body { font-family: Arial, sans-serif; margin: 40px; }\n')
        f.write('h1 { text-align: center; }\n')
        f.write('h2 { margin-top: 30px; }\n')
        f.write('.autor { font-style: italic; }\n')
        f.write('.programa { color: #666; }\n')
        f.write('.divisor { border-bottom: 1px solid #ccc; margin: 20px 0; }\n')
        f.write('</style>\n</head>\n<body>\n')
        
        f.write('<h1>COMPILACIÓN DE CUENTOS</h1>\n')
        f.write('<div class="frase"><em>"A veces, una historia puede ser el puente que necesita un corazón para sanar."</em></div>\n')
        f.write('<div class="mensaje">\n<p>Este libro nace del alma de jóvenes valientes, que se atrevieron a compartir sus emociones, sueños, heridas y esperanzas. Cada cuento aquí escrito es un reflejo auténtico de sus caminos, de sus luchas y de su poder de transformación.</p>\n')
        f.write('<p>Que estas páginas sean no solo un testimonio, sino también una luz para quienes necesitan saber que no están solos.</p>\n')
        f.write('<p>Gracias a cada autor y autora por prestarnos su voz. Gracias por ser una inspiración...</p>\n')
        f.write('<p>Mayo de 2015.</p>\n</div>\n')
        
        f.write('<h2>ÍNDICE</h2>\n<ul>\n')
        for title in index_entries:
            f.write(f'<li><a href="#{title.replace(" ", "_")}">{title}</a></li>\n')
        f.write('</ul>\n')
        
        # Este es un marcador de posición - en una implementación real, se extraería el texto de cada cuento
        for title in index_entries:
            f.write(f'<div id="{title.replace(" ", "_")}" class="cuento">\n')
            f.write(f'<h2>{title}</h2>\n')
            f.write('<p>Contenido del cuento...</p>\n')
            f.write('<div class="divisor"></div>\n')
            f.write(f'<div class="autor">Autor: [Autor del cuento]</div>\n')
            f.write(f'<div class="programa">Programa: [Programa]</div>\n')
            f.write('</div>\n')
        
        f.write('</body>\n</html>')
    
    print("Proceso completado. Se han generado los archivos Compilacion_Cuentos.docx y Compilacion_Cuentos.html")

if __name__ == "__main__":
    main()
